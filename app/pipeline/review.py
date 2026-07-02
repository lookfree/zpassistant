"""阶段4：全局校验（本地预检 + 模型评审）与 Word 渲染导出。"""
import json
from pathlib import Path

from docx import Document
from pydantic import BaseModel

from app.config import settings
from app.zhipu.llm import chat_json
from app.prompts import load_prompt
from app.pipeline import state
from app.pipeline.generate import _iter_sections

EXPORT_DIR = settings.data_dir / "exports"

class ReviewSchema(BaseModel):
    issues: list = []
    summary: str = ""

def full_text(task: dict) -> str:
    parts = [f"# {task['outline'].get('title', task.get('project_name', '技术方案'))}"]
    for ch in task["outline"].get("chapters", []):
        parts.append(f"## {ch['no']} {ch['title']}")
        for sec in ch.get("sections", []):
            parts.append(f"### {sec['no']} {sec['title']}")
            content = task["chapters"].get(sec["no"], {}).get("content", "")
            parts.append(content)
    return "\n\n".join(parts)

def _precheck(task: dict) -> list:
    """本地章节完整性预检：大纲有、正文无 → completeness issue。"""
    issues = []
    for ch, sec in _iter_sections(task["outline"]):
        if not task["chapters"].get(sec["no"], {}).get("content"):
            issues.append({"type": "completeness", "chapter": sec["no"],
                           "desc": f"{sec['no']}《{sec['title']}》未生成内容",
                           "suggestion": "回到分段生成补齐本节"})
    return issues

def run_review(task: dict) -> dict:
    task["stage"] = "reviewing"
    state.save_task(task)
    issues = _precheck(task)
    msgs = [{"role": "system", "content": load_prompt("review", task)},
            {"role": "user", "content":
             f"需求解析结果：\n{json.dumps(task['requirements'], ensure_ascii=False)[:6000]}\n\n"
             f"技术方案全文：\n{full_text(task)[:50000]}"}]
    result = chat_json(msgs, ReviewSchema)
    review_data = {"issues": issues + list(result.issues), "summary": result.summary}
    task["review"] = review_data
    task["stage"] = "reviewed"
    state.save_task(task)
    return review_data

def _add_markdown(doc: Document, text: str) -> None:
    """极简 Markdown → docx：### 标题、- 列表、| 表格 |、普通段落。"""
    rows = []
    def flush_table():
        nonlocal rows
        if not rows:
            return
        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = "Table Grid"
        for i, r in enumerate(rows):
            for j, cell in enumerate(r[:len(table.columns)]):
                table.cell(i, j).text = cell
        rows = []
    for line in text.splitlines():
        s = line.strip()
        if s.startswith("|"):
            cells = [c.strip() for c in s.strip("|").split("|")]
            if not all(set(c) <= set("-: ") for c in cells):   # 跳过分隔行
                rows.append(cells)
            continue
        flush_table()
        if not s:
            continue
        if s.startswith("###"):
            doc.add_heading(s.lstrip("#").strip(), level=3)
        elif s.startswith(("-", "*")):
            doc.add_paragraph(s[1:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(s)
    flush_table()

def export_docx(task: dict) -> Path:
    doc = Document()
    doc.add_heading(task["outline"].get("title", "技术方案"), level=0)
    for ch in task["outline"].get("chapters", []):
        doc.add_heading(f"{ch['no']} {ch['title']}", level=1)
        for sec in ch.get("sections", []):
            doc.add_heading(f"{sec['no']} {sec['title']}", level=2)
            _add_markdown(doc, task["chapters"].get(sec["no"], {}).get("content", ""))
    EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    path = EXPORT_DIR / f"{task['id']}.docx"
    doc.save(str(path))
    task["stage"] = "exported"
    state.save_task(task)
    return path
