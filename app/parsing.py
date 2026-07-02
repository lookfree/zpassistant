"""上传文档 → 纯文本。招标文件在本地解析后喂给模型。"""
import io

from docx import Document
from pypdf import PdfReader

MAX_CHARS = 60000

def _from_docx(content: bytes) -> str:
    doc = Document(io.BytesIO(content))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            parts.append(" | ".join(c for c in cells if c))
    return "\n".join(parts)

def _from_pdf(content: bytes) -> str:
    reader = PdfReader(io.BytesIO(content))
    return "\n".join((page.extract_text() or "") for page in reader.pages)

def extract_text(filename: str, content: bytes) -> str:
    name = filename.lower()
    if name.endswith(".docx"):
        text = _from_docx(content)
    elif name.endswith(".pdf"):
        text = _from_pdf(content)
    elif name.endswith((".txt", ".md")):
        text = content.decode("utf-8", errors="ignore")
    else:
        raise ValueError("不支持的文件类型（支持 docx/pdf/txt/md）")
    if len(text) > MAX_CHARS:
        text = text[:MAX_CHARS] + "\n（文档过长，已截断）"
    return text
