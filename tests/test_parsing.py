import io
import pytest
from docx import Document
from app.parsing import extract_text

def _docx_bytes():
    doc = Document()
    doc.add_paragraph("第一章 总体要求")
    t = doc.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "工期"
    t.rows[0].cells[1].text = "180天"
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

def test_docx_paragraphs_and_tables():
    text = extract_text("a.docx", _docx_bytes())
    assert "第一章 总体要求" in text and "工期" in text and "180天" in text

def test_txt_passthrough():
    assert extract_text("a.txt", "你好".encode()) == "你好"

def test_unsupported_raises():
    with pytest.raises(ValueError):
        extract_text("a.xls", b"")

def test_truncated_over_limit():
    text = extract_text("a.txt", ("字" * 70000).encode())
    assert len(text) < 61000 and "已截断" in text
